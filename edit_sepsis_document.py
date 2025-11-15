from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from copy import deepcopy
import re

def copy_paragraph_format(source_para, target_para):
    """Copy formatting from source paragraph to target"""
    target_para.alignment = source_para.alignment
    target_para.paragraph_format.line_spacing = source_para.paragraph_format.line_spacing
    target_para.paragraph_format.space_before = source_para.paragraph_format.space_before
    target_para.paragraph_format.space_after = source_para.paragraph_format.space_after
    target_para.paragraph_format.left_indent = source_para.paragraph_format.left_indent
    target_para.paragraph_format.right_indent = source_para.paragraph_format.right_indent
    target_para.paragraph_format.first_line_indent = source_para.paragraph_format.first_line_indent

def copy_run_format(source_run, target_run):
    """Copy run formatting"""
    target_run.bold = source_run.bold
    target_run.italic = source_run.italic
    target_run.underline = source_run.underline
    if source_run.font.size:
        target_run.font.size = source_run.font.size
    if source_run.font.name:
        target_run.font.name = source_run.font.name
    if source_run.font.color.rgb:
        target_run.font.color.rgb = source_run.font.color.rgb

def add_citation_to_text(text, citation_nums):
    """Helper to format text with citation"""
    return text, citation_nums

# Load original document
print("Loading original document...")
original_doc = Document('/home/user/articles/Sepsis_withtrials (1).docx')

# Create new edited document
edited_doc = Document()

# Initialize change tracking
changes_log = []
citation_count = 0

print("Processing document with comprehensive edits...")
print("="*80)

# Process each element in the original document
for i, element in enumerate(original_doc.element.body):
    tag = element.tag.split('}')[1] if '}' in element.tag else element.tag
    
    # Handle paragraphs
    if tag == 'p':
        para_index = len([e for e in original_doc.element.body[:i+1] if (e.tag.split('}')[1] if '}' in e.tag else e.tag) == 'p']) - 1
        if para_index < len(original_doc.paragraphs):
            source_para = original_doc.paragraphs[para_index]
            text = source_para.text.strip()
            
            if not text:  # Empty paragraph
                edited_doc.add_paragraph()
                continue
            
            # Create new paragraph in edited document
            new_para = edited_doc.add_paragraph()
            copy_paragraph_format(source_para, new_para)
            
            # Process text with corrections and citations
            # Title
            if text == "Sepsis and Septic Shock":
                run = new_para.add_run(text)
                run.bold = True
                run.font.size = Pt(16)
                new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
            # Section headings (numbered)
            elif re.match(r'^\d+\.\s*[A-Z]', text) or re.match(r'^\d+\.[A-Z]', text):
                for source_run in source_para.runs:
                    new_run = new_para.add_run(source_run.text)
                    copy_run_format(source_run, new_run)
                    
            # Specific content with citations
            elif "Definition and Clinical Significance" in text:
                for source_run in source_para.runs:
                    new_run = new_para.add_run(source_run.text)
                    copy_run_format(source_run, new_run)
                    
            elif "mortality rates ranging from moderate (10%) to substantial (>40%)" in text:
                # Add citation [1]
                new_text = text
                for source_run in source_para.runs:
                    new_run = new_para.add_run(source_run.text)
                    copy_run_format(source_run, new_run)
                # Add citation
                cite_run = new_para.add_run("[1]")
                cite_run.font.superscript = True
                citation_count += 1
                changes_log.append(f"Added citation [1] to mortality statistics")
                
            elif "most recent consensus defines sepsis" in text:
                # Rewrite with citation
                new_text = "The most recent consensus defines sepsis as a life-threatening organ dysfunction caused by a dysregulated host response to infection."
                new_para.add_run(new_text)
                cite_run = new_para.add_run("[1]")
                cite_run.font.superscript = True
                new_text2 = " This definition, established in 2016, represents a fundamental shift from previous classifications by focusing on organ dysfunction rather than inflammatory response criteria."
                new_para.add_run(new_text2)
                cite_run2 = new_para.add_run("[1]")
                cite_run2.font.superscript = True
                citation_count += 2
                changes_log.append(f"Added citations [1] to Sepsis-3 definition")
                
            else:
                # Copy runs with formatting preserved
                for source_run in source_para.runs:
                    new_run = new_para.add_run(source_run.text)
                    copy_run_format(source_run, new_run)
    
    # Handle tables
    elif tag == 'tbl':
        table_index = len([e for e in original_doc.element.body[:i+1] if (e.tag.split('}')[1] if '}' in e.tag else e.tag) == 'tbl']) - 1
        if table_index < len(original_doc.tables):
            source_table = original_doc.tables[table_index]
            
            # Create new table with same dimensions
            new_table = edited_doc.add_table(rows=len(source_table.rows), cols=len(source_table.columns))
            new_table.style = source_table.style
            
            # Copy cell content and formatting
            for row_idx, source_row in enumerate(source_table.rows):
                for col_idx, source_cell in enumerate(source_row.cells):
                    target_cell = new_table.rows[row_idx].cells[col_idx]
                    # Copy cell text and formatting
                    for para in source_cell.paragraphs:
                        if para == source_cell.paragraphs[0]:
                            target_para = target_cell.paragraphs[0]
                        else:
                            target_para = target_cell.add_paragraph()
                        copy_paragraph_format(para, target_para)
                        for run in para.runs:
                            new_run = target_para.add_run(run.text)
                            copy_run_format(run, new_run)

print(f"\nProcessing complete!")
print(f"Citations added: {citation_count}")
print(f"Changes logged: {len(changes_log)}")

# This is a preliminary version - will need full implementation
print("\nNOTE: Creating simplified version for now...")
print("Full detailed editing will be done in complete script...")

